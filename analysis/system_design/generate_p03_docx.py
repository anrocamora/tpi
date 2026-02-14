#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar el documento P03-Diseño de Software con los casos de uso
UC-DS-001 (Subida de datos a landing zone) y UC-DS-002 (Catalogación de ficheros)
para la plataforma TSuPreMe.

Autor: T-Systems
Fecha: 2026-02-14
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Rutas
BASE_PATH = r"C:\Users\A56691985\IdeaProjects\tpi\analysis\system_design"
DIAGRAMS_PATH = os.path.join(BASE_PATH, "diagrams")
OUTPUT_PATH = BASE_PATH


def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level):
    """Añade un encabezado con estilo consistente."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_table(doc, headers, rows, col_widths=None):
    """Crea una tabla con formato profesional."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Encabezados
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E6007E')  # Magenta T-Systems
        for run in header_cells[i].paragraphs[0].runs:
            run.font.color.rgb = None  # Reset color
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Filas
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    
    # Anchos de columna
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    return table


def add_image(doc, image_name, caption, width=6.0):
    """Añade una imagen con título."""
    image_path = os.path.join(DIAGRAMS_PATH, image_name)
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Añadir caption
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Imagen no encontrada: {image_name}]")


def create_document():
    """Crea el documento P03 completo."""
    doc = Document()
    
    # =====================================================================
    # PORTADA
    # =====================================================================
    title = doc.add_heading('P03 - Diseño de Software', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Plataforma TSuPreMe\nSubida y Catalogación de Datos Genómicos')
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Proyecto', 'TSuPreMe - Plataforma de Medicina de Precisión'),
        ('Cliente', 'Servicio Navarro de Salud - Osasunbidea (SNS-O)'),
        ('Versión', '1.0'),
        ('Fecha', '2026-02-14')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # =====================================================================
    # CONTROL DE VERSIONES
    # =====================================================================
    add_heading(doc, 'Control de versiones', 1)
    
    version_headers = ['Versión', 'Fecha', 'Autor', 'Descripción']
    version_rows = [
        ['1.0', '2026-02-14', 'T-Systems', 'Versión inicial con UC-DS-001 y UC-DS-002']
    ]
    add_table(doc, version_headers, version_rows, [2, 3, 4, 8])
    
    doc.add_page_break()
    
    # =====================================================================
    # ÍNDICE (placeholder - Word lo genera automáticamente)
    # =====================================================================
    add_heading(doc, 'Índice', 1)
    doc.add_paragraph('[Actualizar índice en Word: Referencias > Tabla de contenido]')
    doc.add_page_break()
    
    # =====================================================================
    # 1. INTRODUCCIÓN
    # =====================================================================
    add_heading(doc, '1. Introducción', 1)
    
    add_heading(doc, '1.1 Propósito del documento', 2)
    doc.add_paragraph(
        'Este documento constituye la especificación técnica de diseño de software para los '
        'casos de uso de subida y catalogación de datos genómicos en la plataforma TSuPreMe '
        '(T-Systems Precision Medicine). Su objetivo principal es servir como referencia '
        'arquitectónica y guía de implementación para los equipos de desarrollo, operaciones '
        'y calidad.'
    )
    doc.add_paragraph(
        'El documento define en detalle la arquitectura de componentes, los flujos de datos, '
        'los contratos de integración, las decisiones técnicas y los criterios de aceptación '
        'para los casos de uso UC-DS-001 (Subida de datos a landing zone) y UC-DS-002 '
        '(Catalogación de ficheros). Estos dos casos de uso conforman el pipeline inicial '
        'de ingesta de datos genómicos desde los secuenciadores hasta el catálogo centralizado.'
    )
    doc.add_paragraph(
        'Adicionalmente, este documento establece las bases para la trazabilidad técnica, '
        'facilitando la auditoría del sistema y el mantenimiento evolutivo de la plataforma.'
    )
    
    add_heading(doc, '1.2 Alcance', 2)
    doc.add_paragraph(
        'El alcance de este documento cubre el diseño técnico detallado de los dos primeros '
        'casos de uso del dominio de Subida de Datos, que conforman el pipeline de ingesta '
        'de datos genómicos desde los secuenciadores hasta el catálogo centralizado:'
    )
    bullets = [
        'UC-DS-001 - Subida de datos a landing zone: Automatización de la detección de runs '
        'finalizados en unidades de red compartidas y su transferencia confiable al almacenamiento '
        'S3 compatible (THealthLake), incluyendo gestión de eventos, resiliencia y recuperación.',
        'UC-DS-002 - Catalogación de ficheros: Consumo de eventos de completitud y replicación '
        'de la estructura de ficheros en el catálogo Tomic/OpenCGA, con etiquetado semántico '
        'para habilitar búsquedas por run y muestra.',
        'Los casos de uso UC-DS-003 a UC-DS-007 (asociación a peticiones, transformación, '
        'movimiento a ubicación definitiva, creación de casos clínicos y gestión de errores) '
        'quedan referenciados para continuidad de trazabilidad, sin desarrollo técnico en esta versión.'
    ]
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_paragraph(
        'Este documento NO cubre aspectos de infraestructura (despliegue de Kubernetes, '
        'configuración de Kafka, etc.) ni la especificación funcional de usuario, que se '
        'encuentran en documentos complementarios.'
    )

    add_heading(doc, '1.3 Definiciones y acrónimos', 2)
    
    acronym_headers = ['Término', 'Definición']
    acronym_rows = [
        ['TSuPreMe', 'Plataforma de Medicina de Precisión del SNS-O'],
        ['TPI Agent', 'Servicio automatizado de monitorización y subida de datos'],
        ['THealthLake', 'Sistema de almacenamiento S3 compatible para datos de salud'],
        ['TCatalog', 'Sistema de catálogo basado en OpenCGA/Tomic'],
        ['NiFi', 'Apache NiFi - Plataforma de flujo de datos'],
        ['Airflow', 'Apache Airflow - Orquestador de pipelines'],
        ['Kafka', 'Apache Kafka - Plataforma de mensajería distribuida'],
        ['Run', 'Conjunto de datos generado por un secuenciador en una ejecución'],
        ['DLQ', 'Dead Letter Queue - Cola para mensajes no procesables'],
        ['Avro', 'Formato de serialización de datos con esquema'],
        ['JWT', 'JSON Web Token - Token de autenticación'],
        ['S3', 'Simple Storage Service - Almacenamiento de objetos compatible AWS']
    ]
    add_table(doc, acronym_headers, acronym_rows, [4, 13])
    
    doc.add_page_break()
    
    # =====================================================================
    # 2. CATÁLOGO DE CASOS DE USO
    # =====================================================================
    add_heading(doc, '2. Catálogo de casos de uso', 1)
    
    doc.add_paragraph(
        'El dominio de Subida de Datos comprende siete casos de uso que cubren el ciclo '
        'completo de vida de los datos genómicos desde su origen en los secuenciadores '
        'hasta su asociación con casos clínicos. La siguiente tabla presenta el catálogo '
        'completo, indicando el estado de desarrollo de cada uno en este documento.'
    )

    doc.add_paragraph(
        'Los casos UC-DS-001 y UC-DS-002 están completamente desarrollados y constituyen '
        'el núcleo del pipeline de ingesta automatizada. Los demás casos se mantienen '
        'referenciados para garantizar la trazabilidad documental y facilitar su desarrollo '
        'en fases posteriores del proyecto.'
    )
    
    uc_headers = ['Identificador', 'Nombre', 'Estado']
    uc_rows = [
        ['UC-DS-001', 'Subida de datos a landing zone', 'Desarrollado'],
        ['UC-DS-002', 'Catalogación de ficheros', 'Desarrollado'],
        ['UC-DS-003', 'Asociación de ficheros a peticiones', 'Referenciado'],
        ['UC-DS-004', 'Transformación de resultados Nasertic', 'Obsoleto'],
        ['UC-DS-005', 'Movimiento de ficheros a ubicación definitiva', 'Referenciado'],
        ['UC-DS-006', 'Creación de caso clínico/pacientes/muestras', 'Referenciado'],
        ['UC-DS-007', 'Gestión de errores en identificadores de muestra', 'Referenciado']
    ]
    add_table(doc, uc_headers, uc_rows, [3, 10, 4])
    
    doc.add_page_break()
    
    # =====================================================================
    # 3. TRAZABILIDAD FUNCIONAL
    # =====================================================================
    add_heading(doc, '3. Trazabilidad funcional', 1)
    
    doc.add_paragraph(
        'La trazabilidad funcional establece el vínculo entre los casos de uso definidos '
        'en los requisitos y los componentes técnicos que los implementan. Esta trazabilidad '
        'es esencial para verificar que todos los requisitos están cubiertos y para '
        'identificar el impacto de cambios futuros en el sistema.'
    )

    doc.add_paragraph(
        'La siguiente tabla mapea cada caso de uso con su componente principal responsable, '
        'los componentes de soporte que colaboran en su ejecución, y las evidencias técnicas '
        'que demuestran su implementación. Esta información facilita tanto la auditoría '
        'de conformidad como el análisis de impacto ante modificaciones.'
    )
    
    trace_headers = ['Caso de uso', 'Componente principal', 'Componentes de soporte', 'Evidencia técnica']
    trace_rows = [
        ['UC-DS-001', 'TPI Agent Service', 'Kafka Events/State, S3', 'Scheduler + UploadService + UploadStateStore'],
        ['UC-DS-002', 'NiFi Pipeline', 'Airflow DAG, Tomic API, DistributedMapCache', 'DAG + Process Group documentado']
    ]
    add_table(doc, trace_headers, trace_rows, [3, 4, 5, 5])
    
    doc.add_page_break()
    
    # =====================================================================
    # 4. ARQUITECTURA DE REFERENCIA
    # =====================================================================
    add_heading(doc, '4. Arquitectura de referencia', 1)
    
    doc.add_paragraph(
        'La arquitectura de referencia para los casos de uso UC-DS-001 y UC-DS-002 sigue '
        'un modelo event-driven (dirigido por eventos) que desacopla completamente la '
        'fase de subida de datos de la fase de catalogación. Este desacoplamiento proporciona '
        'múltiples beneficios operativos y técnicos:'
    )

    benefits = [
        'Escalabilidad independiente: Cada componente puede escalar según su carga específica '
        'sin afectar a los demás.',
        'Resiliencia: Un fallo en la catalogación no impide que continúen las subidas, y viceversa.',
        'Reprocesamiento: Los eventos almacenados en Kafka permiten reprocesar datos sin '
        'necesidad de volver a subirlos.',
        'Observabilidad: El flujo de eventos proporciona trazabilidad completa de cada operación.',
        'Flexibilidad: Nuevos consumidores pueden suscribirse a los eventos sin modificar '
        'los productores existentes.'
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        'El siguiente diagrama muestra la visión de alto nivel de la arquitectura, '
        'identificando los componentes principales y sus interacciones.'
    )
    
    add_heading(doc, '4.1 Diagrama de contexto', 2)
    
    add_image(doc, 'contexto_uc1_uc2.png', 'Figura 1: Diagrama de contexto UC-DS-001 y UC-DS-002')
    
    doc.add_paragraph()
    add_image(doc, 'diagrama_componentes.png', 'Figura 2: Diagrama de componentes detallado del sistema', 6.5)

    add_heading(doc, '4.2 Descripción de componentes', 2)
    
    comp_headers = ['Componente', 'Tecnología', 'Responsabilidad']
    comp_rows = [
        ['TPI Agent', 'Spring Boot + Java', 'Monitorización de unidad de red, upload a S3, publicación de eventos Kafka'],
        ['Kafka Events', 'Apache Kafka', 'Bus de eventos para comunicación asíncrona entre componentes'],
        ['Kafka State', 'Apache Kafka (compactado)', 'Persistencia de estado de uploads para recuperación'],
        ['THealthLake S3', 'S3 Compatible', 'Almacenamiento de objetos para datos genómicos'],
        ['Airflow DAG', 'Apache Airflow', 'Orquestación del ciclo de vida del pipeline NiFi'],
        ['NiFi Pipeline', 'Apache NiFi', 'Procesamiento de eventos y catalogación en Tomic'],
        ['Tomic API', 'OpenCGA REST API', 'Catálogo de ficheros y metadatos genómicos'],
        ['TCatalog', 'PostgreSQL', 'Base de datos del catálogo']
    ]
    add_table(doc, comp_headers, comp_rows, [3, 4, 10])
    
    add_heading(doc, '4.3 Flujo de datos', 2)
    
    doc.add_paragraph(
        'El flujo de datos describe el recorrido completo de la información desde su origen '
        'en los secuenciadores hasta su disponibilidad en el catálogo para consulta por '
        'usuarios clínicos y analistas. Este flujo está diseñado para ser completamente '
        'automatizado, minimizando la intervención manual y garantizando la trazabilidad.'
    )

    doc.add_paragraph('El flujo sigue la siguiente secuencia de operaciones:')

    flow_steps = [
        'Generación del run: El secuenciador (MiSeq, NextSeq) genera un conjunto de ficheros '
        'de datos crudos y, al finalizar, crea el fichero RunCompletionStatus.xml que indica '
        'que el run está completo y listo para procesamiento.',
        'Detección automática: El TPI Agent Service escanea periódicamente la unidad de red '
        'compartida y detecta los runs completos por la presencia del fichero de completitud.',
        'Transferencia a S3: El agente sube todos los ficheros del run al almacenamiento '
        'THealthLake (S3 compatible), preservando la estructura de directorios original.',
        'Publicación de eventos: Durante la subida, el agente publica eventos Kafka que '
        'permiten el seguimiento en tiempo real del progreso y notifican la completitud.',
        'Orquestación de catalogación: El DAG de Airflow gestiona el ciclo de vida del '
        'pipeline NiFi, arrancándolo cuando hay trabajo y deteniéndolo tras un período de inactividad.',
        'Catalogación en Tomic: El pipeline NiFi consume los eventos de completitud y replica '
        'la estructura de ficheros en el catálogo, añadiendo metadatos y tags de búsqueda.',
        'Disponibilidad para usuarios: Los ficheros catalogados quedan disponibles para '
        'búsqueda en Tomic y acceso directo desde S3 mediante unidades de red montadas.'
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # =====================================================================
    # 5. DIAGRAMA DE CLASES
    # =====================================================================
    add_heading(doc, '5. Diagrama de clases', 1)
    
    doc.add_paragraph(
        'El diagrama de clases presenta las entidades principales del diseño y sus relaciones, '
        'proporcionando una visión estructural del sistema. Este diagrama es fundamental para '
        'comprender cómo se organizan los componentes internamente y cómo colaboran entre sí '
        'para implementar los casos de uso.'
    )

    doc.add_paragraph(
        'El diseño sigue principios de responsabilidad única y separación de concerns, donde '
        'cada clase tiene un propósito bien definido. Las clases se agrupan en dos dominios '
        'principales: el TPI Agent (responsable de la subida) y el Pipeline NiFi (responsable '
        'de la catalogación).'
    )
    
    add_image(doc, 'diagrama_clases_uc1_uc2.png', 'Figura 3: Diagrama de clases UC-DS-001 y UC-DS-002', 6.5)

    add_heading(doc, '5.1 Clases del TPI Agent (UC-DS-001)', 2)
    
    agent_classes = [
        ('DirectoryUploadScheduler', 'Planificador que escanea periódicamente el directorio de entrada '
         'buscando runs completos (fichero RunCompletionStatus.xml)'),
        ('UploadService', 'Servicio principal de subida que gestiona el ciclo de vida del upload, '
         'incluyendo estrategias single-part y multipart'),
        ('UploadStateStore', 'Almacén de estado basado en Kafka compactado para recuperación '
         'tras reinicios del agente'),
        ('UploadEventPublisher', 'Publicador de eventos Kafka (UPLOAD_STARTED, PROGRESS, COMPLETED, FAILED)')
    ]
    for class_name, desc in agent_classes:
        p = doc.add_paragraph()
        p.add_run(f'{class_name}: ').bold = True
        p.add_run(desc)
    
    add_heading(doc, '5.2 Clases del Pipeline NiFi (UC-DS-002)', 2)
    
    nifi_classes = [
        ('AirflowDagController', 'Controlador del DAG que gestiona el ciclo de vida del pipeline NiFi'),
        ('NiFiCatalogPipeline', 'Pipeline que consume eventos y ejecuta la catalogación en Tomic'),
        ('TokenCache', 'Cache distribuida para tokens JWT de autenticación con Tomic'),
        ('TomicApiClient', 'Cliente REST para operaciones de creación de directorios, linkado de ficheros y actualización de tags'),
        ('DLQHandler', 'Manejador de errores que envía mensajes no procesables a la Dead Letter Queue')
    ]
    for class_name, desc in nifi_classes:
        p = doc.add_paragraph()
        p.add_run(f'{class_name}: ').bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # =====================================================================
    # 6. UC-DS-001 — SUBIDA DE DATOS A LANDING ZONE
    # =====================================================================
    add_heading(doc, '6. UC-DS-001 — Subida de datos a landing zone', 1)
    
    doc.add_paragraph(
        'El caso de uso UC-DS-001 implementa la automatización completa de la transferencia '
        'de datos genómicos desde las unidades de red donde los secuenciadores depositan '
        'sus resultados hasta el almacenamiento centralizado en la nube (THealthLake S3). '
        'Este proceso es crítico porque constituye el punto de entrada de todos los datos '
        'genómicos a la plataforma TSuPreMe.'
    )

    doc.add_paragraph(
        'El diseño de este caso de uso prioriza tres aspectos fundamentales: la confiabilidad '
        '(ningún dato debe perderse), la eficiencia (minimizar el tiempo de transferencia) '
        'y la trazabilidad (poder seguir el estado de cada subida en todo momento). Para '
        'lograr estos objetivos, se implementa un servicio daemon que opera de forma continua, '
        'detectando automáticamente nuevos runs y gestionando su transferencia con capacidades '
        'de recuperación ante fallos.'
    )

    # 6.1 Contrato de entrada
    add_heading(doc, '6.1 Contrato de entrada', 2)
    
    doc.add_paragraph().add_run('Identificador: ').bold = True
    doc.paragraphs[-1].add_run('UC-DS-001')

    doc.add_paragraph().add_run('Requerimiento que satisface: ').bold = True
    doc.paragraphs[-1].add_run('REQ-DS-001')

    doc.add_paragraph().add_run('Componente responsable: ').bold = True
    doc.paragraphs[-1].add_run('TPI Agent Service (Spring Boot 3.2.5 + Java 24)')

    doc.add_paragraph().add_run('Trigger de run listo: ').bold = True
    doc.paragraphs[-1].add_run('Presencia del fichero RunCompletionStatus.xml en la raíz del run.')
    
    doc.add_paragraph().add_run('Ubicación monitorizada: ').bold = True
    doc.paragraphs[-1].add_run('Unidad de red compartida (configurable via AGENT_SOURCE_DIR)')

    doc.add_paragraph().add_run('Frecuencia de escaneo: ').bold = True
    doc.paragraphs[-1].add_run('Cada 30 segundos (configurable via AGENT_SCAN_INTERVAL_MS)')

    doc.add_paragraph().add_run('Precondiciones:').bold = True
    preconditions = [
        'El run es legible desde el punto de montaje SMB/NFS del agente',
        'Existe conectividad al endpoint S3 y al broker Kafka',
        'El agente tiene permisos de escritura en S3 y publicación en Kafka',
        'El directorio del run contiene RunCompletionStatus.xml'
    ]
    for pc in preconditions:
        doc.add_paragraph(pc, style='List Bullet')
    
    doc.add_paragraph().add_run('Actores:').bold = True
    actors = [
        'Usuario operador: Coloca runs en la unidad de red compartida',
        'TPI Agent Service: Daemon automatizado de monitorización y subida',
        'Secuenciadores: Illumina MiSeq, NextSeq (generan RunCompletionStatus.xml)'
    ]
    for actor in actors:
        doc.add_paragraph(actor, style='List Bullet')
    
    doc.add_paragraph().add_run('Criterios de validación del run:').bold = True
    validation = [
        'Es un directorio (no fichero suelto)',
        'Contiene RunCompletionStatus.xml en raíz o subcarpetas',
        'No tiene upload activo previo (evitar duplicación)'
    ]
    for v in validation:
        doc.add_paragraph(v, style='List Bullet')

    # 6.2 Flujo principal
    add_heading(doc, '6.2 Flujo principal', 2)
    
    doc.add_paragraph(
        'El flujo principal describe la secuencia completa de operaciones que ejecuta el '
        'TPI Agent desde el momento en que detecta un run finalizado hasta que todos sus '
        'ficheros están almacenados en S3 y el evento de completitud ha sido publicado. '
        'Este flujo está diseñado para ser robusto ante interrupciones y eficiente en el '
        'uso de recursos de red.'
    )

    doc.add_paragraph(
        'El diagrama de secuencia siguiente ilustra las interacciones entre los diferentes '
        'componentes durante una ejecución exitosa del caso de uso (happy path).'
    )
    
    add_image(doc, 'secuencia_uc1_happy_path.png', 'Figura 6: Diagrama de secuencia - Flujo principal UC-DS-001')

    doc.add_paragraph().add_run('Descripción del flujo:').bold = True
    
    flow_desc = [
        ('Detección', 'El agente escanea periódicamente (cada 30 segundos) el directorio configurado '
         'buscando carpetas con el fichero RunCompletionStatus.xml'),
        ('Preparación', 'Una vez detectado un run completo, se mueve a la zona de trabajo (source/) '
         'y se lista recursivamente para calcular los bytes totales'),
        ('Publicación UPLOAD_STARTED', 'Se publica el evento inicial en Kafka incluyendo el catálogo '
         'completo del run (estructura Folder recursiva, ~10 MB)'),
        ('Subida paralela', 'Los ficheros se suben en paralelo usando un pool de 10 hilos configurables. '
         'Por cada fichero completado se publica UPLOAD_PROGRESS (~500 bytes)'),
        ('Finalización', 'Al completar todos los ficheros, se publica UPLOAD_COMPLETED y se mueve '
         'el run a la carpeta completed/')
    ]
    for step, desc in flow_desc:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)
    
    # Añadir estructura de almacenamiento en S3
    doc.add_paragraph()
    doc.add_paragraph().add_run('Estructura de almacenamiento en S3 (Landing Zone):').bold = True
    doc.add_paragraph(
        'Los ficheros se organizan preservando la estructura original del run:'
    )
    doc.add_paragraph('s3://{bucket}/agent/{source_id}/{agent_id}/{run_id}/{path_relativo}', style='List Bullet')
    doc.add_paragraph(
        'Ejemplo: s3://genomica-s3/agent/MiSeq/tsupreme-agent-001/M05089_155_000000000-CT8YM/Data/...'
    )

    # 6.3 Máquina de estados
    add_heading(doc, '6.3 Máquina de estados y resiliencia', 2)
    
    doc.add_paragraph(
        'La gestión del estado de cada upload es fundamental para garantizar la resiliencia '
        'del sistema. El TPI Agent mantiene una máquina de estados que registra la situación '
        'de cada run en proceso, permitiendo recuperar el trabajo tras interrupciones como '
        'reinicios del servicio, cortes de red o fallos de hardware.'
    )

    doc.add_paragraph(
        'El estado de cada upload se persiste en un tópico Kafka compactado, lo que proporciona '
        'una base de datos distribuida y tolerante a fallos sin necesidad de infraestructura '
        'adicional. Cuando el agente arranca, consulta este tópico para identificar uploads '
        'que quedaron pendientes y decide si reanudarlos o marcarlos como abandonados.'
    )

    add_image(doc, 'estados_upload_uc1.png', 'Figura 4: Máquina de estados del upload')
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Vista detallada con notas explicativas:').bold = True
    add_image(doc, 'diagrama_estados_upload.png', 'Figura 5: Estados del upload con anotaciones de comportamiento', 6.5)

    states_headers = ['Estado', 'Descripción', 'Transiciones posibles']
    states_rows = [
        ['DETECTED', 'Run detectado con RunCompletionStatus.xml', 'STARTED'],
        ['STARTED', 'Upload iniciado, evento UPLOAD_STARTED publicado', 'IN_PROGRESS'],
        ['IN_PROGRESS', 'Subida de ficheros en curso', 'COMPLETED, FAILED, ABORTED, ABANDONED'],
        ['COMPLETED', 'Todos los ficheros subidos correctamente', 'Estado final'],
        ['FAILED', 'Error irrecuperable tras agotar reintentos', 'Estado final'],
        ['ABORTED', 'Upload cancelado (fichero eliminado durante subida)', 'Estado final'],
        ['ABANDONED', 'Upload abandonado (más de 24h sin completar)', 'Estado final']
    ]
    add_table(doc, states_headers, states_rows, [3, 7, 7])
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Aspectos de resiliencia:').bold = True
    resilience = [
        'Persistencia de estado en tópico Kafka compactado para recuperación post-restart',
        'Reintentos automáticos en operaciones S3 multipart por parte',
        'Abort explícito de multipart ante fallo irrecuperable',
        'Backoff exponencial entre reintentos (1s → 2s → 4s → 8s)'
    ]
    for r in resilience:
        doc.add_paragraph(r, style='List Bullet')
    
    # 6.4 Reglas de particionado multipart
    add_heading(doc, '6.4 Reglas de particionado multipart', 2)
    
    doc.add_paragraph(
        'La subida de ficheros a S3 utiliza diferentes estrategias según el tamaño del fichero '
        'para optimizar el rendimiento y la resiliencia. Los ficheros grandes se dividen en '
        'partes que se suben independientemente, permitiendo reintentos granulares sin perder '
        'el trabajo ya completado. Esta técnica es especialmente importante para ficheros FASTQ '
        'que pueden superar varios gigabytes.'
    )

    doc.add_paragraph(
        'El siguiente diagrama ilustra el flujo de decisión y las operaciones involucradas '
        'en una subida multipart típica.'
    )

    add_image(doc, 'diagrama_flujo_multipart.png', 'Figura 7: Flujo de subida multipart')

    partition_headers = ['Tamaño del fichero', 'Estrategia', 'Método S3']
    partition_rows = [
        ['0 bytes', 'Single-part especial', 'PutObject con cuerpo vacío'],
        ['< 64 MB (configurable)', 'Single-part', 'PutObject directo'],
        ['≥ 64 MB', 'Multipart', 'CreateMultipartUpload + UploadPart + CompleteMultipartUpload']
    ]
    add_table(doc, partition_headers, partition_rows, [5, 5, 7])
    
    doc.add_paragraph()
    doc.add_paragraph(
        'El umbral de 64 MB es configurable mediante la propiedad agent.upload.part-size-mi-b. '
        'Para ficheros grandes, cada parte se sube independientemente con su propio ETag, '
        'permitiendo reintentos granulares sin repetir partes ya subidas.'
    )
    
    # 6.5 Eventos y contrato Kafka
    add_heading(doc, '6.5 Eventos y contrato Kafka', 2)
    
    doc.add_paragraph(
        'El sistema de eventos Kafka es el mecanismo de comunicación entre el TPI Agent y '
        'los sistemas consumidores, principalmente el pipeline de catalogación (UC-DS-002). '
        'Los eventos se serializan en formato Apache Avro, que proporciona un esquema tipado '
        'y evolución de esquemas compatible.'
    )

    doc.add_paragraph(
        'Se han definido cuatro tipos de eventos funcionales que cubren todo el ciclo de vida '
        'de un upload. La siguiente tabla describe cada tipo, cuándo se publica y qué información '
        'contiene:'
    )

    events_headers = ['Evento', 'Cuándo se publica', 'Incluye catálogo', 'Tamaño aproximado']
    events_rows = [
        ['UPLOAD_STARTED', 'Al iniciar el upload del run', 'Sí', '~10 MB'],
        ['UPLOAD_PROGRESS', 'Por cada fichero completado', 'No', '~500 bytes'],
        ['UPLOAD_COMPLETED', 'Al finalizar todos los ficheros', 'Sí', '~10 MB'],
        ['UPLOAD_FAILED', 'Si falla tras agotar reintentos', 'Sí', '~10 MB']
    ]
    add_table(doc, events_headers, events_rows, [4, 5, 3, 5])
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Contrato mínimo para activar UC-DS-002:').bold = True
    contract = [
        'eventType == "UPLOAD_COMPLETED"',
        'uploadId (identificador único del upload)',
        'agentId (identificador del agente)',
        'folder (estructura recursiva con files[] y folders[])'
    ]
    for c in contract:
        doc.add_paragraph(c, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Tópicos Kafka:').bold = True
    topics = [
        'tpi.uploads.{agent-id}.events.v1: Eventos funcionales (retención 7 días)',
        'tpi.uploads.{agent-id}.state.v1: Estado de uploads (compactado, retención infinita)'
    ]
    for t in topics:
        doc.add_paragraph(t, style='List Bullet')
    
    # Añadir sección de rendimiento
    doc.add_paragraph()
    add_heading(doc, '6.5.1 Rendimiento y optimizaciones', 3)

    doc.add_paragraph().add_run('Subida paralela:').bold = True
    perf_items = [
        '10 hilos concurrentes por defecto (configurable via AGENT_CONCURRENT_UPLOADS)',
        'Run típico de 20 GB / 60.000 ficheros: ~1.5-2 horas',
        'Mejora respecto a subida secuencial: 60-70% más rápido'
    ]
    for p_item in perf_items:
        doc.add_paragraph(p_item, style='List Bullet')

    doc.add_paragraph().add_run('Reducción de tráfico Kafka:').bold = True
    doc.add_paragraph(
        'Solo UPLOAD_STARTED y UPLOAD_COMPLETED incluyen el catálogo completo (~10 MB). '
        'Los eventos UPLOAD_PROGRESS son ligeros (~500 bytes). Para un run con 60.000 ficheros, '
        'esto representa una reducción del 99.99% en tráfico de mensajes.'
    )

    perf_headers = ['Conexión', 'Hilos recomendados']
    perf_rows = [
        ['< 50 Mbps', '5 hilos'],
        ['50-200 Mbps', '10 hilos (default)'],
        ['> 200 Mbps', '15-20 hilos']
    ]
    add_table(doc, perf_headers, perf_rows, [8, 9])

    # 6.6 Riesgos técnicos
    add_heading(doc, '6.6 Riesgos técnicos y mitigaciones', 2)
    
    risks_headers = ['Riesgo', 'Impacto', 'Mitigación']
    risks_rows = [
        ['Corte de red con S3', 'Subida incompleta', 'Retry + resume + estado en Kafka compactado'],
        ['Run corrupto/incompleto', 'Catálogo inconsistente', 'Trigger estricto por RunCompletionStatus.xml'],
        ['Volumen alto de ficheros pequeños', 'Latencia elevada', 'Paralelismo acotado + eventos ligeros en progreso'],
        ['Disco lleno en servidor', 'Subida bloqueada', 'Monitorización de espacio + alertas'],
        ['Credenciales S3 expiradas', 'Fallo de autenticación', 'Rotación automática + alertas de expiración']
    ]
    add_table(doc, risks_headers, risks_rows, [5, 4, 8])
    
    # 6.7 Recuperación tras reinicio
    add_heading(doc, '6.7 Recuperación tras reinicio del agente', 2)
    
    doc.add_paragraph(
        'El sistema garantiza la continuidad operacional sin duplicar cargas ni perder '
        'trazabilidad tras un reinicio inesperado del agente.'
    )
    
    add_image(doc, 'secuencia_recuperacion_uc1_restart.png', 
              'Figura 8: Secuencia de recuperación tras reinicio')

    doc.add_paragraph().add_run('Algoritmo de recuperación:').bold = True
    recovery_steps = [
        'Al arrancar, el agente consulta el UploadStateStore (Kafka compactado)',
        'Se identifican uploads en estado IN_PROGRESS o STARTED',
        'Para cada upload pendiente, se verifica el estado en S3 (listParts/headObject)',
        'Si el upload ya está completo en S3, se publica UPLOAD_COMPLETED (idempotente)',
        'Si hay partes faltantes y el upload tiene menos de 24h, se reanuda',
        'Si el upload tiene más de 24h, se marca como ABANDONED',
        'Si hay error irrecuperable, se aborta el multipart y publica UPLOAD_FAILED'
    ]
    for i, step in enumerate(recovery_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # =====================================================================
    # 7. UC-DS-002 — CATALOGACIÓN DE FICHEROS
    # =====================================================================
    add_heading(doc, '7. UC-DS-002 — Catalogación de ficheros', 1)
    
    doc.add_paragraph(
        'El caso de uso UC-DS-002 implementa la catalogación automática de los ficheros '
        'subidos a S3 en el sistema de catálogo Tomic (basado en OpenCGA). La catalogación '
        'es esencial porque transforma los ficheros almacenados en "datos descubribles": '
        'sin ella, los usuarios no podrían buscar ni localizar los ficheros de interés.'
    )

    doc.add_paragraph(
        'El diseño de este caso de uso sigue el principio de isomorfismo: la estructura de '
        'directorios y ficheros en el catálogo replica exactamente la estructura en S3. '
        'Esto simplifica la navegación y garantiza que cualquier fichero accesible desde S3 '
        'tenga su correspondiente entrada en el catálogo con metadatos enriquecidos.'
    )

    doc.add_paragraph(
        'La catalogación se enriquece con tags semánticos que permiten búsquedas por run '
        '(todos los ficheros de una ejecución de secuenciador) y por muestra (todos los '
        'ficheros asociados a una muestra biológica específica). Este etiquetado facilita '
        'enormemente el trabajo de los analistas y clínicos.'
    )

    # Datos de identificación
    doc.add_paragraph()
    id_data = [
        ('Identificador', 'UC-DS-002'),
        ('Requerimiento que satisface', 'REQ-DS-002'),
        ('Componentes responsables', 'Apache Airflow (DAG) + Apache NiFi (Pipeline)'),
        ('Trigger de activación', 'Evento Kafka UPLOAD_COMPLETED publicado por UC-DS-001')
    ]
    for label, value in id_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_paragraph().add_run('Principios de diseño:').bold = True
    principles = [
        'Isomorfismo de estructura: Los paths en el catálogo replican exactamente los keys en S3, '
        'facilitando la correlación entre ambos sistemas.',
        'Tags semánticos: Cada fichero se etiqueta con identificadores de run y muestra para '
        'habilitar búsquedas eficientes.',
        'Jerarquía completa: Todos los directorios intermedios se crean automáticamente, '
        'respetando las dependencias padre-hijo.',
        'Idempotencia: El reprocesamiento de eventos no causa errores ni duplicados, '
        'tratando respuestas 409 (Conflict) como éxito.'
    ]
    for pr in principles:
        doc.add_paragraph(pr, style='List Bullet')

    # 7.1 Orquestación Airflow
    add_heading(doc, '7.1 Orquestación Airflow', 2)
    
    doc.add_paragraph(
        'La orquestación mediante Apache Airflow separa claramente el gobierno operativo '
        '(cuándo y cómo ejecutar el procesamiento) del procesamiento de datos en sí '
        '(implementado en NiFi). Esta separación proporciona flexibilidad operativa y '
        'visibilidad centralizada de las ejecuciones.'
    )

    doc.add_paragraph(
        'El DAG (Directed Acyclic Graph) de Airflow implementa un patrón de control del '
        'ciclo de vida que arranca el pipeline NiFi, monitorea su actividad, y lo detiene '
        'de forma controlada cuando no hay más trabajo o tras un timeout configurable. '
        'Esto evita el consumo innecesario de recursos cuando no hay eventos que procesar.'
    )
    
    add_image(doc, 'orquestacion_uc2_airflow_nifi.png', 
              'Figura 9: Orquestación Airflow para pipeline NiFi')

    dag_tasks = [
        ('check_nifi_availability', 'Verifica que NiFi responde correctamente'),
        ('start_nifi_pipeline', 'Arranca los Process Groups configurados'),
        ('monitor_nifi_pipeline', 'Monitorea la actividad del pipeline (hasta 10 intentos)'),
        ('wait_before_stop', 'Espera el timeout configurado o indefinidamente'),
        ('stop_nifi_pipeline', 'Detiene los Process Groups'),
        ('trigger_emergency_stop_dag', 'Safety net para garantizar parada de NiFi')
    ]
    
    doc.add_paragraph().add_run('Tareas del DAG:').bold = True
    for task, desc in dag_tasks:
        p = doc.add_paragraph()
        p.add_run(f'{task}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Mecanismos de seguridad multinivel:').bold = True
    safety = [
        'Task explícita stop_nifi_pipeline con trigger_rule=ALL_DONE',
        'Trigger de DAG de emergencia STOP_NIFI_EMERGENCY',
        'Callbacks on_success/on_failure del DAG',
        'Método on_kill() en sensores Wait'
    ]
    for s in safety:
        doc.add_paragraph(s, style='List Bullet')
    
    # 7.2 Diseño funcional pipeline NiFi
    add_heading(doc, '7.2 Diseño funcional de pipeline NiFi', 2)
    
    doc.add_paragraph(
        'El pipeline NiFi implementa la lógica de catalogación mediante una cadena de '
        'procesadores que transforman los eventos Kafka en operaciones sobre la API de Tomic. '
        'NiFi fue elegido por su capacidad de manejar flujos de datos complejos con '
        'trazabilidad incorporada y tolerancia a fallos.'
    )

    doc.add_paragraph(
        'El diseño del pipeline sigue un patrón de procesamiento en fases, donde cada fase '
        'tiene una responsabilidad específica y puede ejecutarse de forma independiente. '
        'La sincronización entre fases se realiza mediante el mecanismo Wait/Notify de NiFi, '
        'que garantiza que los directorios se creen antes de intentar linkear los ficheros.'
    )

    add_image(doc, 'pipeline_nifi_detallado_uc2.png',
              'Figura 10: Pipeline NiFi detallado para catalogación')

    doc.add_paragraph().add_run('Fases del pipeline:').bold = True
    
    pipeline_phases = [
        ('Consumo Kafka', 'ConsumeKafkaRecord lee eventos Avro del tópico de uploads'),
        ('Filtrado', 'Se filtran solo eventos con eventType=UPLOAD_COMPLETED'),
        ('Parseo', 'Se parsea la estructura Folder recursiva para extraer directorios y ficheros'),
        ('Construcción de tareas', 'Se generan listas ordenadas de directorios (padre→hijo) y ficheros'),
        ('Autenticación', 'Se obtiene/renueva token JWT desde cache o mediante login'),
        ('Creación de directorios', 'POST /files/create secuencial para cada directorio'),
        ('Sincronización', 'Wait/Notify asegura que todos los directorios existen antes de linkear'),
        ('Link de ficheros', 'POST /files/link paralelo para cada fichero'),
        ('Actualización de tags', 'POST /files/update para añadir tags run_* y sample_*')
    ]
    for phase, desc in pipeline_phases:
        p = doc.add_paragraph()
        p.add_run(f'{phase}: ').bold = True
        p.add_run(desc)
    
    # 7.3 Secuencia de catalogación
    add_heading(doc, '7.3 Secuencia de catalogación en Tomic', 2)
    
    doc.add_paragraph(
        'La interacción con la API de Tomic sigue una secuencia específica que respeta las '
        'dependencias entre entidades. Los directorios deben existir antes de poder crear '
        'ficheros dentro de ellos, y los ficheros deben existir antes de poder asignarles tags.'
    )

    doc.add_paragraph(
        'El siguiente diagrama de secuencia ilustra las llamadas HTTP a la API de Tomic '
        'durante la catalogación de un run típico, incluyendo la autenticación mediante JWT.'
    )

    add_image(doc, 'secuencia_uc2_catalogacion.png',
              'Figura 11: Secuencia de catalogación en Tomic API')

    doc.add_paragraph().add_run('Operaciones API Tomic:').bold = True
    
    api_ops = [
        ('POST /users/login', 'Obtención de token JWT para autenticación'),
        ('POST /files/create', 'Creación de directorio (type=DIRECTORY)'),
        ('POST /files/link', 'Linkado de fichero con URI S3'),
        ('POST /files/update', 'Actualización de tags (tagsAction=ADD)')
    ]
    for op, desc in api_ops:
        p = doc.add_paragraph()
        p.add_run(f'{op}: ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Extracción de sampleId:').bold = True
    doc.add_paragraph(
        'El identificador de muestra se extrae del nombre del fichero usando el patrón '
        'regex ^(.+?)_S\\d+_ (formato Illumina estándar). Si no hay coincidencia, '
        'se usa el substring antes del primer guión bajo. Ficheros sin patrón reconocible '
        'se catalogan solo con el tag run_*.'
    )
    
    # Tabla de ejemplos de extracción
    sample_headers = ['Nombre fichero', 'Patrón', 'SampleId extraído']
    sample_rows = [
        ['Sample123_S1_L001_R1.fastq.gz', '^(.+?)_S\\d+_', 'Sample123'],
        ['Sample456_aligned.bam', 'Fallback antes de _', 'Sample456'],
        ['README.txt', 'Sin patrón', 'null (solo tag run_*)']
    ]
    add_table(doc, sample_headers, sample_rows, [6, 5, 6])

    # Jerarquía de tags
    doc.add_paragraph()
    doc.add_paragraph().add_run('Jerarquía de tags en catálogo:').bold = True

    tags_hier = [
        ('agent/{source_id}', 'Sin tags'),
        ('agent/{source_id}/{agent_id}', 'Sin tags'),
        ('agent/{source_id}/{agent_id}/{runId}', 'Tags: [run_{runId}]'),
        ('Subcarpetas', 'Tags: [run_{runId}, sample_*, ...] (samples de ficheros descendientes)'),
        ('Ficheros', 'Tags: [run_{runId}] + [sample_{sampleId}] si aplica')
    ]
    for path, tags in tags_hier:
        p = doc.add_paragraph()
        p.add_run(f'{path}: ').bold = True
        p.add_run(tags)

    # 7.4 Idempotencia
    add_heading(doc, '7.4 Idempotencia, concurrencia y orden', 2)
    
    doc.add_paragraph(
        'La idempotencia es un principio de diseño fundamental en este sistema. Dado que '
        'Kafka proporciona semánticas at-least-once (al menos una vez), los eventos pueden '
        'ser procesados más de una vez en caso de reintentos o recuperación de fallos. '
        'Por tanto, todas las operaciones de catalogación están diseñadas para ser seguras '
        'ante ejecuciones repetidas.'
    )

    doc.add_paragraph(
        'El orden de ejecución es igualmente crítico. Los directorios deben crearse antes '
        'que los ficheros que contienen, y este orden debe respetarse estrictamente para '
        'evitar errores de "padre no encontrado" en la API de Tomic.'
    )

    doc.add_paragraph().add_run('Reglas de idempotencia:').bold = True
    idempotence = [
        'Respuestas HTTP 409 (Conflict) se tratan como éxito (recurso ya existe)',
        'tagsAction=ADD no duplica tags existentes',
        'El mismo evento puede reprocesarse sin efectos adversos'
    ]
    for i in idempotence:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_paragraph().add_run('Reglas de orden:').bold = True
    order_rules = [
        'Directorios se crean en orden padre→hijo (ordenados por profundidad de path)',
        'Los ficheros no se linkean hasta que todos sus directorios padres existan',
        'El mecanismo Wait/Notify sincroniza las fases de directorios y ficheros'
    ]
    for o in order_rules:
        doc.add_paragraph(o, style='List Bullet')
    
    doc.add_paragraph().add_run('Concurrencia:').bold = True
    concurrency = [
        'Creación de directorios: Secuencial (por dependencias padre-hijo)',
        'Link de ficheros: Paralelo (Concurrent Tasks configurable)',
        'Actualización de tags: Paralelo (sin dependencias)'
    ]
    for c in concurrency:
        doc.add_paragraph(c, style='List Bullet')
    
    # 7.5 Gestión de errores y DLQ
    add_heading(doc, '7.5 Gestión de errores y DLQ', 2)
    
    doc.add_paragraph(
        'El manejo de errores distingue entre errores transitorios (que pueden resolverse '
        'con reintentos) y errores permanentes (que requieren intervención humana). Esta '
        'distinción es crucial para mantener el sistema funcionando sin bloqueos mientras '
        'se asegura que ningún dato se pierde silenciosamente.'
    )

    doc.add_paragraph(
        'Los errores permanentes se envían a una Dead Letter Queue (DLQ), un tópico Kafka '
        'especial donde los mensajes problemáticos quedan almacenados para análisis posterior. '
        'El equipo de operaciones puede revisar estos mensajes, corregir el problema subyacente '
        'y relanzar el procesamiento de forma controlada.'
    )

    add_image(doc, 'errores_uc2_dlq.png',
              'Figura 13: Gestión de errores y Dead Letter Queue')

    error_headers = ['Código HTTP', 'Tipo', 'Acción']
    error_rows = [
        ['2xx', 'Éxito', 'Continuar al siguiente paso'],
        ['409', 'Conflicto (ya existe)', 'Tratar como éxito (idempotencia)'],
        ['401', 'No autorizado', 'Renovar token JWT y reintentar'],
        ['429', 'Rate limit', 'Reintentar con backoff exponencial'],
        ['5xx', 'Error servidor', 'Reintentar con backoff exponencial'],
        ['4xx (otros)', 'Error cliente', 'Enviar a DLQ (no recuperable)']
    ]
    add_table(doc, error_headers, error_rows, [4, 5, 8])
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Dead Letter Queue (DLQ):').bold = True
    doc.add_paragraph(
        'Los mensajes que no pueden procesarse tras agotar reintentos se envían al tópico '
        'DLQ para análisis manual. El mensaje incluye el payload original y el motivo de rechazo.'
    )
    
    # 7.6 Máquina de estados catalogación
    add_heading(doc, '7.6 Máquina de estados de catalogación', 2)
    
    add_image(doc, 'estado_catalogacion_uc2.png', 
              'Figura 14: Máquina de estados de catalogación')

    cat_states_headers = ['Estado', 'Descripción']
    cat_states_rows = [
        ['EVENT_RECEIVED', 'Evento consumido desde Kafka'],
        ['VALIDATED', 'Evento validado (eventType=UPLOAD_COMPLETED, campos obligatorios presentes)'],
        ['DIRS_CREATED', 'Todos los directorios creados exitosamente en Tomic'],
        ['FILES_LINKED', 'Todos los ficheros linkeados con su URI S3'],
        ['TAGS_UPDATED', 'Tags run_* y sample_* aplicados a todos los ficheros'],
        ['COMPLETED', 'Catalogación completada exitosamente'],
        ['DISCARDED', 'Evento descartado (tipo no relevante)'],
        ['RETRYING', 'En proceso de reintento tras error transitorio'],
        ['DLQ', 'Enviado a Dead Letter Queue por error no recuperable']
    ]
    add_table(doc, cat_states_headers, cat_states_rows, [4, 13])
    
    # 7.7 Secuencia reproceso DLQ
    add_heading(doc, '7.7 Secuencia de reproceso desde DLQ', 2)
    
    add_image(doc, 'secuencia_reproceso_dlq_uc2.png', 
              'Figura 15: Secuencia de reproceso desde DLQ')

    doc.add_paragraph().add_run('Procedimiento operativo de reproceso:').bold = True
    reprocess_steps = [
        'El equipo SRE identifica mensajes en DLQ por uploadId/runId',
        'Se analiza la causa raíz del fallo (logs, métricas)',
        'Se corrige el problema subyacente si es necesario',
        'Se lanza un job de replay controlado que consume mensajes de DLQ',
        'Por cada mensaje, se reinyecta al pipeline NiFi',
        'Si el reproceso es exitoso, se confirma el ACK',
        'Si persiste el error, se mantiene en DLQ con información enriquecida',
        'Se genera reporte de reproceso con pendientes'
    ]
    for i, step in enumerate(reprocess_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # Añadir sección de rendimiento y timeline
    add_heading(doc, '7.8 Rendimiento y timeline típico', 2)

    doc.add_paragraph().add_run('Catalogación típica (run con 100 ficheros):').bold = True

    cat_perf_headers = ['Fase', 'Tiempo esperado']
    cat_perf_rows = [
        ['Consumo Kafka + Construcción tareas', '< 5 s'],
        ['Login Tomic (si no cacheado)', '< 500 ms'],
        ['Creación directorios (secuencial)', '3-10 dirs × 200 ms = 0.6-2 s'],
        ['Sincronización Wait/Notify', '< 1 s'],
        ['Link ficheros (paralelo)', '100 ficheros ÷ N hilos × 300 ms'],
        ['Update tags (paralelo)', '100 ficheros ÷ N hilos × 200 ms'],
        ['Total', '< 2 minutos']
    ]
    add_table(doc, cat_perf_headers, cat_perf_rows, [8, 9])

    doc.add_paragraph()
    doc.add_paragraph().add_run('Timeline end-to-end típico (Run 5 GB, 100 ficheros):').bold = True

    timeline_headers = ['Tiempo', 'Evento', 'Componente']
    timeline_rows = [
        ['T+0s', 'Secuenciador completa, genera RunCompletionStatus.xml', 'Secuenciador'],
        ['T+30s', 'Agente detecta run completo en siguiente escaneo', 'TPI Agent'],
        ['T+31s', 'Publicado UPLOAD_STARTED en Kafka', 'TPI Agent'],
        ['T+32s - T+10m', 'Subida paralela de ficheros a S3', 'TPI Agent'],
        ['T+10m', 'Publicado UPLOAD_COMPLETED en Kafka', 'TPI Agent'],
        ['T+10m+1s', 'NiFi consume evento y parsea Folder', 'NiFi Pipeline'],
        ['T+10m+5s', 'Login Tomic + creación directorios', 'NiFi Pipeline'],
        ['T+10m+10s', 'Link de ficheros + update tags', 'NiFi Pipeline'],
        ['T+12m', '✅ Catalogación completa', 'Sistema']
    ]
    add_table(doc, timeline_headers, timeline_rows, [3, 9, 5])

    # Diagrama de flujo completo end-to-end
    doc.add_paragraph()
    add_heading(doc, '7.9 Flujo completo end-to-end', 2)

    doc.add_paragraph(
        'El siguiente diagrama muestra la secuencia completa de operaciones desde que el '
        'secuenciador genera un run hasta que los ficheros están catalogados y disponibles '
        'para consulta.'
    )

    add_image(doc, 'diagrama_flujo_completo.png',
              'Figura 12: Flujo completo end-to-end UC-DS-001 y UC-DS-002', 6.5)

    doc.add_page_break()
    
    # =====================================================================
    # 8. ARQUITECTURA DE DATOS
    # =====================================================================
    add_heading(doc, '8. Arquitectura de datos', 1)
    
    doc.add_paragraph(
        'La arquitectura de datos define cómo fluye la información a través del sistema, '
        'qué formatos se utilizan y cómo se organizan los datos en cada almacén. Un diseño '
        'coherente de la arquitectura de datos es esencial para garantizar la interoperabilidad '
        'entre componentes y la consistencia de la información.'
    )

    add_heading(doc, '8.1 Flujo de datos end-to-end', 2)
    
    doc.add_paragraph(
        'El siguiente diagrama muestra el recorrido completo de los datos desde su origen '
        '(unidad de red donde el secuenciador deposita los ficheros) hasta su destino final '
        '(catálogo TCatalog donde quedan disponibles para búsqueda). El flujo atraviesa '
        'múltiples sistemas, cada uno con su responsabilidad específica.'
    )

    add_image(doc, 'arquitectura_datos_uc1_uc2.png',
              'Figura 16: Arquitectura de datos UC-DS-001 y UC-DS-002')

    doc.add_paragraph()
    add_image(doc, 'diagrama_usuarios_y_almacenamiento.png',
              'Figura 17: Interacción de usuarios y sistemas de almacenamiento', 6.5)

    doc.add_paragraph(
        'El flujo de datos atraviesa múltiples sistemas desde el origen (unidad de red) '
        'hasta el destino final (catálogo TCatalog), con almacenamiento intermedio en S3 '
        'y comunicación asíncrona via Kafka.'
    )
    
    # Estructura de almacenamiento detallada
    add_heading(doc, '8.1.1 Estructura de almacenamiento en THealthLake', 3)

    doc.add_paragraph().add_run('Landing Zone (zona temporal):').bold = True
    doc.add_paragraph('Ubicación: agent/')
    doc.add_paragraph(
        'Estructura: agent/{source_id}/{agent_id}/{run_id}/{...ficheros...}'
    )
    doc.add_paragraph(
        'Los ficheros mantienen su estructura original del run, preservando la jerarquía '
        'de carpetas exacta. El fichero RunCompletionStatus.xml se ignora en la subida.'
    )

    doc.add_paragraph().add_run('Almacenamiento definitivo (futuro UC-DS-005):').bold = True
    doc.add_paragraph('Ubicación: data/')
    doc.add_paragraph(
        'Estructura objetivo: data/sample/{sample_id}/rawdata/ para datos crudos, '
        'data/sample/{sample_id}/results/ para resultados de análisis.'
    )

    add_heading(doc, '8.2 Modelo de eventos Avro', 2)
    
    doc.add_paragraph(
        'Apache Avro es el formato de serialización elegido para los eventos Kafka debido '
        'a sus ventajas en entornos de integración: esquema tipado que previene errores, '
        'serialización binaria eficiente que reduce el tamaño de los mensajes, y soporte '
        'para evolución de esquemas que facilita el mantenimiento a largo plazo.'
    )

    doc.add_paragraph(
        'El modelo de eventos se estructura en torno a la clase UploadEvent, que contiene '
        'toda la información necesaria para que los consumidores procesen el evento sin '
        'necesidad de consultar sistemas externos. La estructura Folder incluida en el evento '
        'representa el catálogo completo del run de forma recursiva.'
    )

    add_image(doc, 'clases_evento_upload_avro_uc1_uc2.png',
              'Figura 18: Modelo de clases del evento UploadEvent (Avro)')

    doc.add_paragraph().add_run('Estructura del evento UploadEvent:').bold = True
    
    event_fields = [
        ('eventType', 'string', 'Tipo de evento (UPLOAD_STARTED, PROGRESS, COMPLETED, FAILED)'),
        ('uploadId', 'string', 'Identificador único del upload (UUID)'),
        ('agentId', 'string', 'Identificador del agente que realiza la subida'),
        ('runId', 'string', 'Identificador del run (nombre de carpeta)'),
        ('timestamp', 'datetime', 'Marca temporal del evento'),
        ('folder', 'Folder', 'Estructura recursiva con ficheros y subcarpetas'),
        ('metadata', 'map<string,string>', 'Metadatos adicionales opcionales')
    ]
    
    field_headers = ['Campo', 'Tipo', 'Descripción']
    field_rows = [[f[0], f[1], f[2]] for f in event_fields]
    add_table(doc, field_headers, field_rows, [3, 4, 10])
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Estructura del objeto Folder:').bold = True
    
    folder_fields = [
        ('name', 'string', 'Nombre del directorio'),
        ('url', 'string', 'URL S3 del directorio'),
        ('files', 'FileRef[]', 'Lista de ficheros en el directorio'),
        ('folders', 'Folder[]', 'Lista de subdirectorios (recursivo)')
    ]
    
    folder_headers = ['Campo', 'Tipo', 'Descripción']
    folder_rows = [[f[0], f[1], f[2]] for f in folder_fields]
    add_table(doc, folder_headers, folder_rows, [3, 4, 10])
    
    add_heading(doc, '8.3 Claves de correlación', 2)
    
    doc.add_paragraph(
        'Las claves de correlación permiten trazar un dato a través de todo el sistema, '
        'desde su origen hasta su destino final. Son fundamentales para el diagnóstico de '
        'problemas, la auditoría y la generación de informes operativos.'
    )

    doc.add_paragraph(
        'Las siguientes claves están presentes en todos los eventos y registros del sistema, '
        'permitiendo correlacionar información entre componentes:'
    )

    correlation = [
        ('uploadId', 'Unidad técnica de correlación (UUID único por upload)'),
        ('runId', 'Unidad funcional de correlación (nombre del run)'),
        ('agentId', 'Identificador del origen de los datos')
    ]
    for key, desc in correlation:
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(desc)
    
    add_heading(doc, '8.4 Taxonomía de tags en catálogo', 2)
    
    doc.add_paragraph(
        'Los tags son etiquetas de metadatos que se aplican a ficheros y directorios en el '
        'catálogo para habilitar búsquedas semánticas. A diferencia de los paths, que son '
        'estructurales, los tags permiten clasificar los datos según criterios de negocio '
        'como el run de origen o la muestra biológica asociada.'
    )

    doc.add_paragraph(
        'La siguiente tabla describe los tipos de tags utilizados, su formato y cuándo se aplican:'
    )

    tags_headers = ['Tag', 'Formato', 'Aplicación']
    tags_rows = [
        ['run_*', 'run_{runId}', 'Todos los directorios y ficheros del run'],
        ['sample_*', 'sample_{sampleId}', 'Ficheros con sampleId extraído del nombre'],
        ['source_id', 'Identificador de origen', 'Trazabilidad de procedencia'],
        ['agent_id', 'Identificador de agente', 'Trazabilidad operativa']
    ]
    add_table(doc, tags_headers, tags_rows, [3, 4, 10])
    
    # Búsquedas en catálogo
    add_heading(doc, '8.5 Búsquedas habilitadas en catálogo', 2)

    doc.add_paragraph('El sistema de tags permite las siguientes búsquedas desde Tomic/TCatalog:')

    search_headers = ['Tipo de búsqueda', 'Ejemplo de tag/query', 'Resultado']
    search_rows = [
        ['Por run', 'run_M05089_155_000000000-CT8YM', 'Todos los ficheros/dirs del run'],
        ['Por muestra', 'sample_S1365399822', 'Todos los ficheros de esa muestra'],
        ['Por path (wildcard)', 'agent/MiSeq/tsupreme-agent-001/M05089*', 'Coincidencias por patrón'],
        ['Combinada', 'run_* AND sample_*', 'Ficheros con ambos criterios']
    ]
    add_table(doc, search_headers, search_rows, [4, 6, 7])

    doc.add_page_break()
    
    # =====================================================================
    # 9. OPERACIÓN, OBSERVABILIDAD Y SRE
    # =====================================================================
    add_heading(doc, '9. Operación, observabilidad y SRE', 1)
    
    doc.add_paragraph(
        'La observabilidad es la capacidad de entender el estado interno de un sistema '
        'a partir de sus salidas externas (logs, métricas, trazas). Un sistema observable '
        'permite detectar problemas rápidamente, diagnosticar sus causas y verificar que '
        'las soluciones aplicadas son efectivas.'
    )

    doc.add_paragraph(
        'El diseño de UC-DS-001 y UC-DS-002 incorpora observabilidad desde el inicio, '
        'no como una funcionalidad añadida posteriormente. Cada componente genera telemetría '
        'que alimenta dashboards operativos y sistemas de alertas, permitiendo una operación '
        'proactiva según las prácticas de Site Reliability Engineering (SRE).'
    )

    add_image(doc, 'observabilidad_sre_uc1_uc2.png',
              'Figura 19: Arquitectura de observabilidad')

    add_heading(doc, '9.1 Métricas recomendadas', 2)
    
    doc.add_paragraph().add_run('UC-DS-001 (Subida):').bold = True
    metrics_uc1 = [
        'Tiempo total de subida por run',
        'Throughput en MB/s',
        'Ratio de fallos por fichero',
        'Número de reintentos por upload',
        'Tamaño medio de run procesado'
    ]
    for m in metrics_uc1:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_paragraph().add_run('UC-DS-002 (Catalogación):').bold = True
    metrics_uc2 = [
        'Lag del consumidor Kafka',
        'Tiempo de catalogación por run',
        'Tasa de mensajes enviados a DLQ',
        'Latencia de operaciones API Tomic',
        'Ratio de renovaciones de token JWT'
    ]
    for m in metrics_uc2:
        doc.add_paragraph(m, style='List Bullet')
    
    add_heading(doc, '9.2 Alertas recomendadas', 2)
    
    alerts_headers = ['Alerta', 'Condición', 'Severidad']
    alerts_rows = [
        ['Upload sin completar', 'Run detectado sin UPLOAD_COMPLETED en ventana esperada', 'Warning'],
        ['DLQ creciente', 'Incremento continuo de mensajes en DLQ', 'Critical'],
        ['Fallos de autenticación', 'Token/login failures repetidos en Tomic', 'Critical'],
        ['Lag Kafka elevado', 'Consumer lag > umbral durante periodo sostenido', 'Warning'],
        ['Error rate elevado', 'Tasa de errores HTTP > 5% en ventana', 'Warning']
    ]
    add_table(doc, alerts_headers, alerts_rows, [5, 8, 4])
    
    add_heading(doc, '9.3 Dashboards recomendados', 2)
    
    dashboards = [
        ('Dashboard UC-DS-001', 'Estado de uploads activos, throughput, errores, runs en cola'),
        ('Dashboard UC-DS-002', 'Eventos procesados, lag Kafka, tasa DLQ, latencias API'),
        ('Dashboard de salud', 'Disponibilidad de componentes, conectividad, recursos')
    ]
    for name, content in dashboards:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(content)
    
    doc.add_page_break()
    
    # =====================================================================
    # 10. SEGURIDAD E INTEGRIDAD DE DATOS
    # =====================================================================
    add_heading(doc, '10. Seguridad e integridad de datos', 1)
    
    doc.add_paragraph(
        'La seguridad e integridad de los datos genómicos es un requisito fundamental del sistema. '
        'Los datos genómicos son información sensible de salud que debe protegerse durante '
        'la transmisión y el almacenamiento, cumpliendo con normativas como el RGPD.'
    )

    doc.add_paragraph(
        'El diseño implementa múltiples capas de protección que garantizan tres propiedades '
        'esenciales: confidencialidad (los datos solo son accesibles por usuarios autorizados), '
        'integridad (los datos no se alteran durante la transmisión) y disponibilidad '
        '(los datos permanecen accesibles cuando se necesitan).'
    )
    
    add_heading(doc, '10.1 Integridad y completitud de datos', 2)
    
    doc.add_paragraph().add_run('Transmisión hacia S3 (UC-DS-001):').bold = True
    integrity_s3 = [
        'Validación por checksum MD5 calculado antes de subir y verificado por S3',
        'Verificación de Content-Length para detectar transmisiones truncadas',
        'Reintentos automáticos con backoff exponencial ante fallos de red',
        'Verificación post-upload mediante headObject comparando tamaño'
    ]
    for i in integrity_s3:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_paragraph().add_run('Mensajería con Kafka (UC-DS-001 y UC-DS-002):').bold = True
    integrity_kafka = [
        'Configuración acks=all para garantizar replicación completa',
        'Productores idempotentes (enable.idempotence=true) para evitar duplicados',
        'Consumer offsets confirmados solo tras procesamiento exitoso',
        'Tópico de estado compactado como base de datos para recuperación'
    ]
    for i in integrity_kafka:
        doc.add_paragraph(i, style='List Bullet')
    
    add_heading(doc, '10.2 Encriptación en la transmisión', 2)
    
    doc.add_paragraph().add_run('Hacia S3:').bold = True
    encryption_s3 = [
        'TLS 1.2+ automático en todas las comunicaciones (HTTPS)',
        'Server-Side Encryption AES-256 para datos en reposo',
        'Endpoints HTTPS forzados en configuración del cliente'
    ]
    for e in encryption_s3:
        doc.add_paragraph(e, style='List Bullet')
    
    doc.add_paragraph().add_run('Hacia Kafka:').bold = True
    encryption_kafka = [
        'SSL/TLS para conexiones entre productores, brokers y consumidores',
        'Autenticación SASL/SSL opcional para verificar identidad',
        'Certificados gestionados mediante keystores JKS'
    ]
    for e in encryption_kafka:
        doc.add_paragraph(e, style='List Bullet')
    
    add_heading(doc, '10.3 Resumen de garantías', 2)
    
    guarantees_headers = ['Aspecto', 'Mecanismo', 'Resultado']
    guarantees_rows = [
        ['Integridad de archivos', 'MD5 checksum + verificación de tamaño', 'Archivos completos y sin corrupción'],
        ['Completitud de runs', 'Flag RunCompletionStatus.xml + catálogo', 'Solo se suben runs completos'],
        ['Integridad de mensajes', 'Kafka acks=all + idempotencia', 'Sin pérdida ni duplicación'],
        ['Encriptación S3', 'HTTPS/TLS 1.2+ automático', 'Datos cifrados en tránsito'],
        ['Encriptación Kafka', 'SSL/TLS 1.2+ configurado', 'Mensajes cifrados en tránsito'],
        ['Encriptación en reposo', 'S3 SSE-AES256', 'Datos cifrados almacenados'],
        ['Recuperación de fallos', 'Reintentos + reanudación', 'Resiliencia ante interrupciones']
    ]
    add_table(doc, guarantees_headers, guarantees_rows, [5, 6, 6])
    
    add_heading(doc, '10.4 Verificación de integridad', 2)
    
    doc.add_paragraph('Procedimientos para verificar la transmisión correcta de un run:')
    
    verification_steps = [
        'Verificar tamaño en S3 mediante listado recursivo y comparar con origen',
        'Buscar evento UPLOAD_COMPLETED en Kafka con totalBytesUploaded correcto',
        'Verificar que el catálogo folder lista todos los archivos subidos',
        'Confirmar que el run se movió al directorio completed/',
        'Revisar logs del agente buscando confirmación de éxito'
    ]
    for i, step in enumerate(verification_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # =====================================================================
    # 11. CRITERIOS DE ACEPTACIÓN TÉCNICOS
    # =====================================================================
    add_heading(doc, '11. Criterios de aceptación técnicos', 1)
    
    doc.add_paragraph(
        'Los criterios de aceptación técnicos definen las condiciones verificables que '
        'deben cumplirse para considerar que los casos de uso están correctamente implementados. '
        'Estos criterios complementan los criterios de aceptación funcionales definidos en '
        'las especificaciones de requisitos, centrándose en aspectos técnicos y de rendimiento.'
    )

    doc.add_paragraph(
        'Cada criterio está diseñado para ser medible y verificable mediante pruebas automatizadas '
        'o procedimientos de validación definidos. El cumplimiento de estos criterios es '
        'requisito para la aceptación del sistema en cada entorno (DEV, PRE, PRO).'
    )

    doc.add_paragraph()
    doc.add_paragraph().add_run('UC-DS-001 (Subida de datos):').bold = True

    criteria_uc1 = [
        ('CA-01', 'Run detectado en ≤ 30 segundos tras aparecer RunCompletionStatus.xml'),
        ('CA-02', 'Todos los ficheros subidos con paths idénticos al origen'),
        ('CA-03', 'Eventos Kafka UPLOAD_STARTED, UPLOAD_PROGRESS, UPLOAD_COMPLETED publicados correctamente'),
        ('CA-04', 'Upload reanudable tras reinicio del agente (< 24h)'),
        ('CA-05', 'Run de 20 GB / 60k ficheros completa en < 2.5 horas'),
        ('CA-06', 'Ficheros de 0 bytes se suben correctamente usando PutObject')
    ]
    
    criteria_uc1_headers = ['ID', 'Criterio de aceptación']
    criteria_uc1_rows = [[c[0], c[1]] for c in criteria_uc1]
    add_table(doc, criteria_uc1_headers, criteria_uc1_rows, [2, 15])

    doc.add_paragraph()
    doc.add_paragraph().add_run('UC-DS-002 (Catalogación de ficheros):').bold = True

    criteria_uc2 = [
        ('CA-07', 'Paths en Tomic idénticos a keys en S3'),
        ('CA-08', 'Directorios creados en orden padre→hijo'),
        ('CA-09', 'Ficheros no se linkean hasta que directorios existen (sincronización)'),
        ('CA-10', 'Tags run_* y sample_* aplicados correctamente'),
        ('CA-11', 'Reprocesar evento no causa errores (idempotencia con 409)'),
        ('CA-12', 'Run con 100 ficheros cataloga en < 2 minutos')
    ]

    criteria_uc2_headers = ['ID', 'Criterio de aceptación']
    criteria_uc2_rows = [[c[0], c[1]] for c in criteria_uc2]
    add_table(doc, criteria_uc2_headers, criteria_uc2_rows, [2, 15])

    doc.add_page_break()
    
    # =====================================================================
    # 12. DECISIONES DE DISEÑO (ADR-LITE)
    # =====================================================================
    add_heading(doc, '12. Decisiones de diseño (ADR-lite)', 1)
    
    doc.add_paragraph(
        'Las decisiones de diseño documentan las elecciones arquitectónicas significativas '
        'realizadas durante el proyecto, incluyendo el contexto que las motivó, las alternativas '
        'consideradas y las consecuencias esperadas. Esta documentación es valiosa para '
        'entender el "por qué" detrás del diseño actual y facilitar futuras evoluciones.'
    )

    doc.add_paragraph(
        'Se utiliza un formato ADR-lite (Architecture Decision Record simplificado) que '
        'captura la esencia de cada decisión sin la sobrecarga documental de formatos más '
        'extensos. Cada decisión incluye contexto, la decisión tomada y sus consecuencias '
        'tanto positivas como negativas.'
    )
    
    decisions = [
        {
            'id': 'ADR-001',
            'title': 'Arquitectura event-driven entre UC-DS-001 y UC-DS-002',
            'context': 'Se necesita desacoplar la subida de datos de la catalogación para permitir '
                       'escalabilidad y resiliencia independientes.',
            'decision': 'Usar Apache Kafka como bus de eventos asíncrono entre el TPI Agent y el '
                        'pipeline NiFi de catalogación.',
            'consequences': [
                'Positivo: Desacoplamiento temporal - subida y catalogación operan independientemente',
                'Positivo: Buffer natural ante picos de carga',
                'Positivo: Reprocesamiento posible desde eventos almacenados',
                'Negativo: Complejidad operativa adicional (gestión de Kafka)'
            ]
        },
        {
            'id': 'ADR-002',
            'title': 'Airflow para control del ciclo de vida de NiFi',
            'context': 'El pipeline NiFi necesita arrancar/parar de forma controlada, evitando '
                       'consumo de recursos cuando no hay trabajo.',
            'decision': 'Usar Apache Airflow como orquestador del ciclo de vida del pipeline NiFi, '
                        'separando el control operativo del procesamiento de datos.',
            'consequences': [
                'Positivo: Control centralizado del ciclo de vida',
                'Positivo: Visibilidad de ejecuciones en UI de Airflow',
                'Positivo: Múltiples mecanismos de seguridad para garantizar parada',
                'Negativo: Dependencia adicional de Airflow'
            ]
        },
        {
            'id': 'ADR-003',
            'title': 'Persistencia de estado de upload en Kafka compactado',
            'context': 'El agente puede reiniciarse inesperadamente y necesita reanudar uploads '
                       'sin perder progreso ni duplicar trabajo.',
            'decision': 'Usar un tópico Kafka compactado como almacén de estado distribuido, '
                        'permitiendo recuperación del estado de uploads tras reinicio.',
            'consequences': [
                'Positivo: Recuperación automática sin base de datos externa',
                'Positivo: Estado siempre sincronizado con eventos publicados',
                'Positivo: Retención infinita de última versión de cada upload',
                'Negativo: Latencia de recuperación proporcional a número de uploads activos'
            ]
        },
        {
            'id': 'ADR-004',
            'title': 'Catalogación idempotente como criterio de robustez',
            'context': 'Los eventos Kafka pueden reprocesarse (at-least-once semantics) y la '
                       'catalogación debe ser tolerante a duplicados.',
            'decision': 'Diseñar todas las operaciones de catalogación como idempotentes, '
                        'tratando respuestas 409 (Conflict) como éxito.',
            'consequences': [
                'Positivo: Tolerancia a reprocesamiento de eventos',
                'Positivo: Simplificación de lógica de error handling',
                'Positivo: Operaciones de retry seguras',
                'Negativo: Imposibilidad de detectar verdaderos conflictos de negocio'
            ]
        }
    ]
    
    for dec in decisions:
        add_heading(doc, f'{dec["id"]}: {dec["title"]}', 2)
        
        p = doc.add_paragraph()
        p.add_run('Contexto: ').bold = True
        p.add_run(dec['context'])
        
        p = doc.add_paragraph()
        p.add_run('Decisión: ').bold = True
        p.add_run(dec['decision'])
        
        p = doc.add_paragraph()
        p.add_run('Consecuencias:').bold = True
        for cons in dec['consequences']:
            doc.add_paragraph(cons, style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # =====================================================================
    # ANEXO: MODELO DE PROCESOS
    # =====================================================================
    add_heading(doc, 'Anexo A: Modelo de procesos', 1)

    doc.add_paragraph(
        'Este anexo documenta los procesos de negocio asociados a los casos de uso UC-DS-001 '
        'y UC-DS-002. Los procesos describen el flujo de actividades desde una perspectiva '
        'operativa, complementando la visión técnica presentada en las secciones anteriores.'
    )

    doc.add_paragraph(
        'Cada proceso define su objetivo, el flujo de pasos y las reglas de negocio que '
        'deben respetarse durante su ejecución. Esta información es útil para el personal '
        'de operaciones y para la definición de procedimientos de soporte.'
    )

    add_heading(doc, 'A.1 PR-DS-001: Transferencia de datos a landing zone', 2)

    doc.add_paragraph().add_run('Descripción: ').bold = True
    doc.paragraphs[-1].add_run(
        'Proceso encargado de la transferencia automatizada de datos de secuenciación '
        'desde las carpetas o unidades de red designadas a la landing zone en THealthLake.'
    )

    doc.add_paragraph().add_run('Flujo del proceso:').bold = True
    pr001_steps = [
        'Monitorización de la carpeta o unidad de red designada (escaneo cada 30s)',
        'Detección de nueva carpeta/run completo (verificación de RunCompletionStatus.xml)',
        'Movimiento a zona de trabajo (source/)',
        'Inicio de transferencia a S3 (multipart para ficheros grandes, paralelo)',
        'Publicación de eventos Kafka (UPLOAD_STARTED, UPLOAD_PROGRESS, UPLOAD_COMPLETED)',
        'Persistencia de estado en Kafka state topic (compactado)',
        'Movimiento a carpeta completed/ tras éxito'
    ]
    for i, step in enumerate(pr001_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph().add_run('Reglas de negocio:').bold = True
    pr001_rules = [
        'La transferencia solo se inicia cuando se detecta RunCompletionStatus.xml',
        'Debe garantizarse integridad de datos (checksums, ETags)',
        'Eventos Kafka para cada etapa relevante (trazabilidad completa)',
        'Reanudación automática en caso de interrupción (< 24h)',
        'Ficheros > 64 MB usan multipart upload'
    ]
    for rule in pr001_rules:
        doc.add_paragraph(rule, style='List Bullet')

    add_heading(doc, 'A.2 PR-DS-002: Catalogación inicial de ficheros', 2)

    doc.add_paragraph().add_run('Descripción: ').bold = True
    doc.paragraphs[-1].add_run(
        'Proceso encargado de crear entidades en TCatalog para runs/carpetas y ficheros subidos, '
        'incluyendo metadatos para trazabilidad y búsqueda.'
    )

    doc.add_paragraph().add_run('Flujo del proceso:').bold = True
    pr002_steps = [
        'Login en TCatalog para obtención de token JWT',
        'Consumo de evento UPLOAD_COMPLETED desde Kafka',
        'Parseo de estructura folder recursiva',
        'Construcción de listas de directorios (ordenados por profundidad) y ficheros',
        'Creación secuencial de directorios (POST /files/create)',
        'Sincronización Wait/Notify antes de crear ficheros',
        'Creación paralela de ficheros (POST /files/link)',
        'Aplicación de tags run_* y sample_* (POST /files/update)'
    ]
    for i, step in enumerate(pr002_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph().add_run('Reglas de negocio:').bold = True
    pr002_rules = [
        'El identificador de muestra se extrae del nombre del fichero según patrón Illumina',
        'Directorios se crean antes que ficheros (sincronización obligatoria)',
        'Conflictos 409 se tratan como éxito (idempotencia)',
        'Token JWT se cachea y renueva automáticamente',
        'Errores no recuperables se envían a DLQ para análisis manual'
    ]
    for rule in pr002_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # ANEXO: REQUISITOS FUNCIONALES
    # =====================================================================
    add_heading(doc, 'Anexo B: Requisitos funcionales', 1)

    doc.add_paragraph(
        'Este anexo presenta los requisitos funcionales derivados del análisis de los casos '
        'de uso UC-DS-001 y UC-DS-002. Cada requisito está identificado de forma única, '
        'descrito con detalle suficiente para su implementación y priorizado según su '
        'criticidad para el funcionamiento del sistema.'
    )

    doc.add_paragraph(
        'Los requisitos con prioridad ALTA son obligatorios para la puesta en producción. '
        'Los requisitos con prioridad MEDIA pueden implementarse en fases posteriores sin '
        'comprometer la funcionalidad básica del sistema.'
    )

    requirements = [
        ('RF-001', 'Monitorización automatizada',
         'El sistema debe monitorizar automáticamente directorios compartidos detectando runs completos '
         '(presencia de RunCompletionStatus.xml). Frecuencia configurable, por defecto 30 segundos.', 'ALTA'),
        ('RF-002', 'Subida fiable a S3',
         'El sistema debe transferir runs completos a S3 THealthLake preservando estructura recursiva. '
         'Estrategias: single-part (< 64 MB), multipart (≥ 64 MB). Reintentos hasta 3 intentos con backoff.', 'ALTA'),
        ('RF-003', 'Trazabilidad por eventos',
         'Cada operación significativa debe generar evento Kafka trazable. Eventos obligatorios: '
         'UPLOAD_STARTED, UPLOAD_PROGRESS, UPLOAD_COMPLETED, UPLOAD_FAILED.', 'ALTA'),
        ('RF-004', 'Orquestación de pipeline',
         'Airflow DAG debe gestionar ciclo de vida completo de pipeline NiFi: inicio, monitorización, '
         'parada y safety net de emergencia.', 'ALTA'),
        ('RF-005', 'Catalogación isomorfa',
         'El catálogo Tomic debe replicar exactamente la estructura S3. Regla: path_tomic == s3_key.', 'ALTA'),
        ('RF-006', 'Extracción de metadatos',
         'El sistema debe extraer automáticamente sampleId de nombres de ficheros usando patrón '
         'regex ^(.+?)_S\\d+_ o fallback.', 'MEDIA'),
        ('RF-007', 'Etiquetado semántico',
         'Ficheros y directorios deben etiquetarse para búsqueda: run_{runId}, sample_{sampleId}.', 'ALTA'),
        ('RF-008', 'Idempotencia',
         'Reprocesar eventos no debe causar errores. 409 Conflict tratado como éxito.', 'ALTA'),
        ('RF-009', 'Reintentos inteligentes',
         'Errores transitorios (5xx, 429, timeouts) reintentados automáticamente. '
         'Errores permanentes (4xx) enviados a DLQ.', 'ALTA'),
        ('RF-010', 'Gestión de token JWT',
         'El sistema debe minimizar logins reutilizando tokens JWT. Cache distribuido con renovación automática.', 'MEDIA'),
        ('RF-011', 'Sincronización de fases',
         'Ficheros no deben linkearse hasta que sus directorios padres existan (Wait/Notify).', 'ALTA'),
        ('RF-012', 'Rendimiento paralelo',
         'El sistema debe maximizar throughput: 10 hilos para upload, ficheros en paralelo para catalogación.', 'MEDIA')
    ]

    req_headers = ['ID', 'Nombre', 'Descripción', 'Prioridad']
    req_rows = [[r[0], r[1], r[2], r[3]] for r in requirements]
    add_table(doc, req_headers, req_rows, [2, 4, 9, 2])

    doc.add_page_break()

    # =====================================================================
    # ANEXO: DIAGRAMA DE DESPLIEGUE
    # =====================================================================
    add_heading(doc, 'Anexo C: Diagrama de despliegue', 1)

    doc.add_paragraph(
        'El diagrama de despliegue muestra la distribución física de los componentes del '
        'sistema en la infraestructura de ejecución. Este diagrama es útil para el equipo '
        'de operaciones y para planificar la capacidad y alta disponibilidad del sistema.'
    )

    doc.add_paragraph(
        'Los componentes se despliegan en un entorno Kubernetes, aprovechando las capacidades '
        'de orquestación de contenedores para gestión automática de escalado, recuperación '
        'de fallos y actualizaciones progresivas.'
    )

    add_image(doc, 'diagrama_despliegue.png', 
              'Figura C1: Diagrama de despliegue de componentes', 6.5)

    # =====================================================================
    # GUARDAR DOCUMENTO
    # =====================================================================
    output_file = os.path.join(OUTPUT_PATH, 'p03-diseno de software.docx')
    doc.save(output_file)
    print(f"Documento generado exitosamente: {output_file}")
    return output_file


if __name__ == '__main__':
    create_document()

